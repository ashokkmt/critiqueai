from flask import Flask, render_template, request, session, jsonify
import os
import json
import uuid
import google.generativeai as genai
from datetime import datetime, timezone, timedelta
import concurrent.futures
import firebase_admin
from firebase_admin import credentials, storage, firestore
import base64
import PIL.Image
import docx
from odfdo import Document

# Initialize Firebase
cred = credentials.Certificate("code/firebase.json")
firebase_admin.initialize_app(cred, {
    "storageBucket": "instant-theater-449913-h4.firebasestorage.app"
})

bucket = storage.bucket()
db = firestore.client()

with open('code/config.json', 'r') as c:
    params = json.load(c)["params"]

# Configure the Google Generative AI API
genai.configure(api_key=params['gen_api'])
model = genai.GenerativeModel("gemini-2.0-flash")
model_pro = genai.GenerativeModel("gemini-2.0-flash")

# Initialize Flask app
app = Flask(__name__)
app.secret_key = params['session_key']

# Configure upload folder and file size limits
app.config['UPLOAD_FOLDER'] = params['upload_folder']
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # Max file size: 16 MB

pdf_extract_prompt = """Extract all possible information from the given image while maintaining its original structure and meaning. Ensure that:
- **For Text:** Extract all visible text exactly as it appears, preserving formatting, font styles (bold, italic, underline), and special symbols.
- **For Tables:** Identify and extract all tabular data while preserving row and column structure.
- **For Charts & Graphs:** Describe the type of chart, key trends, axes labels, legends, and any numerical values present.
- **For Diagrams & Figures:** Identify objects, flowcharts, and connections, and describe their meaning.
- **For Mathematical Equations:** Extract equations exactly as shown, maintaining proper mathematical notation.
- **For Images with Embedded References:** Identify references, labels, and any related annotations.

If any one of the field is not present then simply write NONE there.
"""

# Define allowed file extensions
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'txt', 'pdf', 'docx', 'odt'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def upload_files(file):
    session_id = session.get('session_id')
    unique_filename = f"{uuid.uuid4()}.{file.filename.split('.')[-1]}"
    blob = bucket.blob(f"sessions/{session_id}/{unique_filename}")
    blob.upload_from_file(file)
    blob.make_public()
    url = f"gs://instant-theater-449913-h4.firebasestorage.app/sessions/{session_id}/{unique_filename}"

    db.collection("user_files").document(unique_filename).set({
        "session_id": session_id,
        "filename": unique_filename,
        "url": url,
        "uploaded_at": datetime.now(timezone.utc)
    })

def get_user_files(session_id):
    docs = db.collection("user_files").where(
        "session_id", "==", session_id).stream()
    file_urls = [doc.to_dict()["url"] for doc in docs]
    return file_urls

def read_file_from_gcs(gs_url):
    file_path = '/'.join(gs_url.split('/')[3:])
    blob = bucket.blob(file_path)
    return blob.download_as_bytes(), blob.name

def get_evaluate(text):
    response = model.generate_content(
        [text, pdf_extract_prompt], generation_config=genai.GenerationConfig(
            max_output_tokens=1000,
            temperature=0.5,))
    return response

def process_txt_file(path):
    with open(path, "r") as text_file:
        doc_text = text_file.read()
    return get_evaluate(doc_text)

def process_pdf_file(path):
    with open(path, "rb") as doc_file:
        doc_data = base64.standard_b64encode(doc_file.read()).decode("utf-8")
    response = model.generate_content(
        [{'mime_type': 'application/pdf', 'data': doc_data}, pdf_extract_prompt], generation_config=genai.GenerationConfig(
            max_output_tokens=200,
            temperature=0.5,))
    return response

def process_img_file(path):
    img_file = PIL.Image.open(path)
    return get_evaluate(img_file)

def process_docx_file(path):
    doc = docx.Document(path)
    full_text = [paragraph.text for paragraph in doc.paragraphs]
    docx_text = '\n'.join(full_text)
    return get_evaluate(docx_text)

def process_odt_file(path):
    odt_file = Document(path)
    text_content = [
        para.text for para in odt_file.body.get_elements("//text:p")]
    odt_text = '\n'.join(text_content)
    return get_evaluate(odt_text)

def process_file(gs_url):
    file_bytes, file_name = read_file_from_gcs(gs_url)
    file_ext = file_name.split(".")[-1].lower()

    if file_ext in ["txt", "csv", "json"]:
        return file_bytes.decode("utf-8")
    elif file_ext == "pdf":
        return process_pdf_file(file_bytes)
    elif file_ext == "docx":
        return process_docx_file(file_bytes)
    elif file_ext in ["png", "jpg", "jpeg"]:
        return process_img_file(file_bytes)
    elif file_ext == "odt":
        return process_odt_file(file_bytes)
    else:
        return f"[Unsupported file type: {file_name}]"

def generate_content_for_user(user_id):
    file_urls = get_user_files(user_id)
    with concurrent.futures.ThreadPoolExecutor() as executor:
        file_contents = list(executor.map(process_file, file_urls))
    combined_text = "\n\n".join(file_contents)
    return combined_text

@app.before_request
def set_session():
    if 'session_id' not in session:
        session['session_id'] = str(uuid.uuid4())

@app.route('/')
def home():
    return render_template("summary.html")

@app.route('/summary')
def summary():
    return render_template("summary.html")

@app.route('/summary_out', methods=['POST'])
def summary_out():
    if request.method == 'POST':
        check_file = 'file' in request.files and request.files['file'].filename
        check_fname = 'fname' in request.form and request.form['fname']
        if check_file and check_fname:
            data = request.form['fname']
            return render_template("summary_out.html", output=f"Your file and text both will get evaluated. Data = {data}")

        elif check_file:
            files = request.files.getlist('file')
            for f in files:
                upload_files(f)
            session_id = session.get('session_id')
            response = generate_content_for_user(session_id)
            return render_template("summary_out.html", output=f"File is generated\n\n{response}")

        elif check_fname:
            data = request.form['fname']
            return render_template("summary_out.html", output=data)

        else:
            return render_template("summary_out.html", output="Invalid input received.")

@app.route('/delete-user-files', methods=['GET'])
def accountcleanup():
    expiration_time = datetime.now(timezone.utc) - timedelta(minutes=5)

    docs = db.collection("user_files").where(
        "uploaded_at", "<", expiration_time).stream()

    deleted_files = []

    for doc in docs:
        data = doc.to_dict()
        session_id = data.get("session_id")
        filename = data.get("filename")

        if session_id and filename:
            file_path = f"sessions/{session_id}/{filename}"
            blob = bucket.blob(file_path)

            try:
                blob.delete()
                db.collection("user_files").document(doc.id).delete()
                deleted_files.append(filename)
            except Exception as e:
                print(f"Failed to delete {file_path}: {e}")

    if not deleted_files:
        return "No Files at Firebase\n", 200

    return jsonify({"deleted_files": deleted_files}), 200

if __name__ == '__main__':
    app.run(debug=True)
