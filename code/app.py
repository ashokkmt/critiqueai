from flask import Flask, render_template, request, session, jsonify, redirect, url_for
import pymupdf
import pdfplumber
import os
import json
import docx
import markdown
from werkzeug.utils import secure_filename
from google import genai
from google.genai import types
from datetime import datetime, timezone, timedelta
from google.cloud import storage, vision, secretmanager
from PIL import Image
from io import BytesIO
import zipfile
import firebase_admin
from firebase_admin import credentials, storage, firestore
import uuid
from google.oauth2 import service_account
import io
import base64


sec_client = secretmanager.SecretManagerServiceClient()

# Firebase Certificate
secret_name = "projects/952301619936/secrets/firebase-config/versions/latest"
response = sec_client.access_secret_version(request={"name": secret_name})
firebase_json = json.loads(response.payload.data.decode("UTF-8"))

# Gemini API key 1
gen_secret = "projects/952301619936/secrets/gemini-api/versions/latest"
get_key = sec_client.access_secret_version(request={"name": gen_secret})
gen_key = get_key.payload.data.decode("UTF-8")

# Gemini API key 2
gen_secret2 = "projects/952301619936/secrets/critiqueai_gen_key/versions/latest"
get_key2 = sec_client.access_secret_version(request={"name": gen_secret2})
gen_key2 = get_key2.payload.data.decode("UTF-8")

# GCS Service Account Vision Certificate
serve_account = "projects/952301619936/secrets/service_account/versions/latest"
get_serve_account = sec_client.access_secret_version(request={"name": serve_account})
service_account_json = json.loads(get_serve_account.payload.data.decode("UTF-8"))


# Updated prompts for AI evaluation
EVALUATION_PROMPT = """Evaluate the following question and answer pair for accuracy and relevance. Provide a concise summary (max 60-70 words, human -like legal language but simple), justification for your evaluations, and suggest specific improvements. Do not include any introductory text. Do not mention the question and answer JUST GIVE EVALUATION AS YOU TOLD.
eg-   Q:
      A:
       
question:
answer: 
"""

SCORE_PROMPT = "Evaluate the following question and answer pair and give it a combined score out of 0 to 10. Just give one word score like 'Score: 7'. (Always give 0 if the answer is absolutely wrong)"


SUMMARY_PROMPT = '''Generate a structured summary of the provided data. The data consists of multiple files, each separated by a line of asterisks (*************************).
DONT ADD ANY SUMMARY STARTING LINE. JUST START SUMMART DIRECTLY.

For each file:
Extract key points and organize them under a relevant heading based on the file's content.
Maintain the summary length in proportion to the content’s significance; do not unnecessarily shorten it.
Use bullet points for clarity and readability.
Do not include any greetings or salutations.

IMP: Format the output for easy readability.
'''

# New prompt for roadmap generation
ROADMAP_PROMPT = '''Generate a concise roadmap for learning {topic} over 4-6 weeks. The roadmap should be divided into logical sections and include:

Key Concepts – Outline 3-5 foundational concepts or principles in {topic}, each explained in 15-20 words.
Practice Tasks – Provide 2-3 hands-on exercises or projects that reinforce the concepts. Each task should have clear instructions and expected outcomes.
Best Resources – List 3-5 effective resources (books, courses, websites, tools) with clickable links and a one-line description for each.
Format the output for easy readability.'''

IMG_PROMPT = "Describe the content of the image extracted from the PDF."

CONTENT_PROMPT = """
You are a professional educator. Generate well-structured study notes for the topic: '{topic}'. 
The student is in '{academic_level}' {course} and prefers '{note_level}' notes. 
Format the notes in '{format_preference}' style. 
Exam-focused: '{exam_focus}'. 
If the topic is technical, include code snippets: '{technical_content}'. 
Urgency level: '{urgency_level}'. 
Adjust content depth, complexity, and length based on these inputs.

Don't any greeting and thank you note at the start or the end.
"""


# Initialize Flask app
app = Flask(__name__)
app.secret_key = os.urandom(24)

cred = credentials.Certificate(firebase_json)
firebase_admin.initialize_app(cred, {
    "storageBucket": "instant-theater-449913-h4.firebasestorage.app"
})


# Configure the Google Generative AI API
client = genai.Client(api_key=gen_key2)
FLASH = 'gemini-2.0-flash'
FLASH_LITE = 'gemini-2.0-flash-lite'

bucket = storage.bucket()
db = firestore.client()

# Initialize GCS client
credentials = service_account.Credentials.from_service_account_info(service_account_json)

# Initialize all Google Cloud clients with the same credentials
vision_client = vision.ImageAnnotatorClient(credentials=credentials)


# Define allowed file extensions
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'txt', 'pdf', 'docx'}


# Function to check allowed file extensions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Function for summary page
def get_evaluation(text, isInstruction=None):
    if isInstruction:  # If an instruction is provided, append it
        full_prompt = f"{SUMMARY_PROMPT}\n\nAdditional Instructions: {isInstruction}"
    else:
        full_prompt = SUMMARY_PROMPT
        
    response = client.models.generate_content(
        model=FLASH,
        contents=[text, full_prompt],
        config= types.GenerateContentConfig(
            max_output_tokens=1000,  # Set the token limit back to 1500
            temperature=0.5)
    )
    
    content = response.text if response.candidates else "No summary generated"
    return content

# <----------------- Evaluation Page Here Starts From Here ------------------->
# Functions for evaluation page
def get_evaluate(text):
    response = client.models.generate_content(
        model=FLASH,
        contents=[text, EVALUATION_PROMPT], 
        config=types.GenerateContentConfig(
            max_output_tokens=1000,
            temperature=0.5))
    score = client.models.generate_content(
        model=FLASH,
        contents=[text, SCORE_PROMPT], 
        config=types.GenerateContentConfig(
            max_output_tokens=1000,
            temperature=0.5))
    return response, score


# Text File
def process_txt_file(content):
    return get_evaluate(content)


# PDF File
def process_pdf_file(pdf_bytes):
    response = client.models.generate_content(
        model=FLASH,
        contents=[types.Part.from_bytes(
            data=pdf_bytes,
            mime_type='application/pdf'),
            EVALUATION_PROMPT], 
        config=types.GenerateContentConfig(
            max_output_tokens=200,
            temperature=0.5,
        )
    )
    score = client.models.generate_content(
        model=FLASH,
        contents=[types.Part.from_bytes(
            data=pdf_bytes,
            mime_type='application/pdf'),
            SCORE_PROMPT], 
        config=types.GenerateContentConfig(
            max_output_tokens=200,
            temperature=0.5,
        )
    )
    return response, score

# Image File
def process_img_file(image_content):
    image = vision.Image(content=image_content)
    text_response = vision_client.text_detection(image=image)
    text = text_response.text_annotations[0].description if text_response.text_annotations else ""
 
    return get_evaluate(text)
    
    
# Docx File
def process_docx_file(doc_bytes):
    doc = docx.Document(BytesIO(doc_bytes)) 
    full_text = [paragraph.text for paragraph in doc.paragraphs]
    docx_text = '\n'.join(full_text)

    return get_evaluate(docx_text) 
        
# <----------------- Evaluate Page Functions Ends Here ------------------>


# <----------------- Helper functions for summary_out page ------------------>
# Image Extraction for PDF
def describe_image_with_gemini(image, flag=False):
    """
    Process image using Gemini with proper encoding and fallback
    Returns: Gemini description (preferred) or Vision API analysis (fallback)
    """

    # Encode to base64
    img_base64 = encode_image(image)

    # Generate description with Gemini
    response = client.models.generate_content(
        model=FLASH_LITE,
        contents=[
            {"text": IMG_PROMPT},
            {"inline_data": {
                "mime_type": "image/png",
                "data": img_base64
            }}
        ],
        config=types.GenerateContentConfig(
            temperature=0.3, 
            max_output_tokens=1000 
        )
    )

    # Return formatted response
    if response.candidates:
        return response.text.strip()
    return "No description generated"

        
# Encoding image for PDF
def encode_image(image):
    """Convert PIL Image to Base64 format for Gemini."""
    if isinstance(image, bytes):
        image = Image.open(io.BytesIO(image))
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    return base64.b64encode(buffered.getvalue()).decode("utf-8")

# Image Extraction for PDF
def extract_images(pdf_file, page_index):
    """Extracts images from a PDF page with Gemini descriptions."""
    page = pdf_file.load_page(page_index)
    image_list = page.get_images(full=True)
    images = []

    for img_index, img in enumerate(image_list, start=1):
        try:
            xref = img[0]
            base_image = pdf_file.extract_image(xref)
            image_bytes = base_image["image"]  # Raw bytes
            image_ext = base_image["ext"]
            image_name = f"image_{page_index+1}_{img_index}.{image_ext}"

            image = Image.open(io.BytesIO(image_bytes))

            # Generate description using Gemini with raw bytes
            image_description = describe_image_with_gemini(image)

            images.append({
                "name": image_name,
                "description": image_description
            })
        except Exception as e:
            images.append(f"[Image extraction error: {str(e)}]")
    return images

# Text Extraction for PDF
def extract_text(pdf_file, page_index):
    """Extracts text from a given PDF page."""
    page = pdf_file.load_page(page_index)
    return page.get_text("text")

# Tables Extraction for PDF
def extract_tables(content, page_index):
    """Extracts tables using pdfplumber."""
    tables = []
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            if page_index < len(pdf.pages):
                page = pdf.pages[page_index]
                extracted_tables = page.extract_tables()
            for table_index, table in enumerate(extracted_tables, start=1):
                tables.append({
                    "table_number": table_index,
                    "data": table
                })
    except Exception as e:
        tables.append(f"[Table error: {str(e)}]")
    return tables

# Process PDF for Page-Wise Extraction
def extract_pdf_content(pdf_bytes):
    """Extracts text, images, and tables from PDF bytes and returns as a string."""
    pdf_file = pymupdf.open(
        stream=pdf_bytes, filetype="pdf")  # Open from bytes
    extracted_data = []

    for page_index in range(len(pdf_file)):
        page_content = {
            "page_number": page_index + 1,
            "text": extract_text(pdf_file, page_index),
            "images": extract_images(pdf_file, page_index),
            # Directly pass bytes
            "tables": extract_tables(pdf_bytes, page_index)
        }
        extracted_data.append(page_content)

    # json_filename = "output_pdf.json"
    # with open(json_filename, "w", encoding="utf-8") as json_file:
    #     json.dump(extracted_data, json_file, indent=4, ensure_ascii=False)

    # print(f"[+] Extraction complete. Data saved to {json_filename}")

    # Convert the extracted_data list to a string
    result = "PDF CONTENT BELOW HERE\n"
    for page in extracted_data:
        result += f"PAGE {page['page_number']}:\n"
        result += f"TEXT:\n{page['text']}\n"

        result += "IMAGES:\n"
        for img in page['images']:
            if isinstance(img, dict):
                result += f"--> {img['name']}: {img['description']}\n"
            else:
                result += f"--> {img}\n"  # Error message case
        result += "\n"

        result += "TABLES:\n"
        for table in page['tables']:
            if isinstance(table, dict):
                result += f"Table {table['table_number']}:\n"
                for row in table['data']:
                    result += f"{row}\n"
                result += "\n"
            else:
                result += f"--> {table}\n"  # Error message case
        result += "---------------------------------------------------\n"

    return result
# PDF FUNCTION ENDS HERE


# Functions to read .docx files
def read_docx_file(docx_content):
    """Read the content of a DOCX file and return it as a string, including tables."""
    doc = docx.Document(BytesIO(docx_content))
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            full_text.append('\t'.join(row_data))

    return '\n'.join(full_text)


def extract_docx_content(content):
    try:
        # Extract regular text content
        text_content = read_docx_file(content)

        # Extract and process images
        image_texts = []
        with zipfile.ZipFile(BytesIO(content)) as zip_file:
            for file_info in zip_file.infolist():
                if file_info.filename.startswith('word/media/'):
                    image_data = zip_file.read(file_info)
                    try:
                        img_text = describe_image_with_gemini(image_data)
                        image_texts.append(img_text)
                    except Exception as e:
                        # Handle image processing errors silently or log them
                        pass

        # Combine all text elements
        combined_text = f"{text_content}\n{' '.join(image_texts)}"
        
        return combined_text

    except Exception as e:
        return f"DOCX Error: {str(e)}"


# Function to upload file to Google Cloud Storage
def upload_files(file):
    session_id = session.get('session_id')
    unique_filename = f"{uuid.uuid4()}.{file.filename.split('.')[-1]}"
    blob = bucket.blob(f"sessions/{session_id}/{unique_filename}")
    blob.upload_from_file(file)
    blob.make_public()
    url = f"gs://instant-theater-449913-h4.firebasestorage.app/sessions/{session_id}/{unique_filename}"

    print(session_id)
    print("file uploaded now updating db")

    db.collection("user_files").document(unique_filename).set({
        "session_id": session_id,
        "filename": unique_filename,
        "url": url,
        "uploaded_at": datetime.now(timezone.utc)
    })

# Function to get file content from Google Cloud Storage
def get_user_files(session_id, time):
    """Fetches gs:// URLs for a user from Firestore."""
    
    docs = db.collection("user_files").where(
        "session_id", "==", session_id).stream()
    file_urls = []
    for doc in docs:
        file_data = doc.to_dict()
        if file_data["uploaded_at"] >= time:
            file_url = file_data["url"]
            file_urls.append(file_url)

    return file_urls


def read_file_from_gcs(gs_url):
    """Reads a file as bytes directly from Firebase Storage."""
    file_path = '/'.join(gs_url.split('/')[3:])
    blob = bucket.blob(file_path)
    print("BLOB NAME - ", blob.name)
    return blob.download_as_bytes(), blob.name


def process_file(gs_url):
    """Processes a file based on its type (Text, PDF, DOCX, Image)."""
    file_bytes, filename = read_file_from_gcs(gs_url)
    file_ext = filename.split(".")[-1].lower()

    # Plain text files
    if file_ext in ["txt", "csv", "json"]:
        return file_bytes.decode("utf-8")

    # PDF files (Text + Images)
    elif file_ext == "pdf":
        final_extracted_content = extract_pdf_content(file_bytes)
        return final_extracted_content

    # Image processing
    elif file_ext in ["png", "jpg", "jpeg"]:
        return describe_image_with_gemini(file_bytes, True)

    # Docx Processing
    elif file_ext == "docx":
        return extract_docx_content(file_bytes)
    else:
        raise ValueError("Unsupported file type")
# <---------- Summary Page Functions Ends Here ------------>


# Set session before each request
@app.before_request
def set_session():
    if 'session_id' not in session:
        session['session_id'] = str(uuid.uuid4())
        print("Session ID - ", session['session_id'])


# Route for home page to load index.html
@app.route('/')
def home():
    if session.get("result_generated"):
        session.pop("result_generated", None)
        print("Result generated marked false")
    return render_template("index.html")


# Route for roadmap page
@app.route('/roadmap')
def roadmap():
    if session.get("result_generated"):
        session.pop("result_generated", None)
        print("Result generated marked false")
    return render_template("roadmap.html")


# Route for summary page
@app.route('/summary')
def summary():
    if session.get("result_generated"):
        session.pop("result_generated", None)
        print("Result generated marked false")
    return render_template("summary.html")


# Route for input page
@app.route('/input')
def input():
    if session.get("result_generated"):
        session.pop("result_generated", None)
        print("Result generated marked false")
    return render_template("input.html")

@app.route('/get-content')
def content():
    if session.get("result_generated"):
        session.pop("result_generated", None)
        print("Result generated marked false")
    return render_template("content.html")

@app.route('/content_out', methods=['POST'])
def generate_content():
    try:
        if request.method == 'POST':
            if session.get("result_generated"):
                print("Rechecked result generation reload")
                return redirect(url_for("get-content"))
            data = request.get_json()
            topic = data.get('topic')
            academic_level = data.get('academic_level')
            course = f"({data.get('course')})" if academic_level == "College" else ""
            note_level = data.get('note_level')
            format_preference = data.get('format_preference')
            technical_content = data.get('technical_content')
            urgency_level = data.get('urgency_level')
            exam_focus = data.get('exam_focus')

            # Format the prompt
            full_prompt = CONTENT_PROMPT.format(
                topic=topic,
                academic_level=academic_level,
                course=course,
                note_level=note_level,
                format_preference=format_preference,
                technical_content=technical_content,
                urgency_level=urgency_level,
                exam_focus=exam_focus
            )
            
            if not data:
                return "Please enter a valid topic", 400
            
            try:
                # full_prompt = CONTENT_PROMPT.format(topic=data)
                # print(full_prompt)
                response = client.models.generate_content(
                    model=FLASH,
                    contents=full_prompt,
                    config=types.GenerateContentConfig(
                        max_output_tokens=8192,
                        temperature=0.5
                    )
                )
                temp = markdown.markdown(response.text)
                session["result_generated"] = True 
                print("result marked true")
                return render_template("content_out.html", output=temp)

            except Exception as e:
                return f"Error generating roadmap: {str(e)}", 500
    except Exception as e:
        print("🔥 ERROR:", str(e))
        return jsonify({"error": str(e)}), 500
    

# New route for generating roadmap
@app.route('/get_roadmap', methods=['POST'])
def get_roadmap():
    if request.method == 'POST':
        if session.get("result_generated"):
            print("Rechecked result generation reload")
            return redirect(url_for("roadmap"))
        data = request.form['topic']

        if not data:
            return "Please enter a valid topic", 400

        try:
            full_prompt = ROADMAP_PROMPT.format(topic=data)
            response = client.models.generate_content(
                model=FLASH,
                contents=full_prompt,
                config=types.GenerateContentConfig(
                    max_output_tokens=2500,
                    temperature=0.5
                )
            )
            temp = markdown.markdown(response.text)
            session["result_generated"] = True 
            print("result marked true")
            return temp if response.candidates else "Failed to generate roadmap"

        except Exception as e:
            return f"Error generating roadmap: {str(e)}", 500

# Route for Summary output page
@app.route('/summary_out', methods=['POST'])
def summary_out():
    if request.method == 'POST':
        if session.get("result_generated"):
            print("Session key 'result_generated' found. Redirecting to /summary...")
            return redirect(url_for("summary"))
        check_file = 'file' in request.files and request.files.getlist('file')
        check_fname = 'fname' in request.form and request.form['fname']
        combined_text = ""
        time = datetime.now(timezone.utc)
        if check_file and check_fname:
            instruction = request.form['fname']
            files = request.files.getlist('file')

            for f in files:
                if f and allowed_file(f.filename):
                    upload_files(f)
                else:
                    combined_text += f"Invalid or unsupported file: {f.filename}\n"
            
            session_id = session.get('session_id')
            file_url = get_user_files(session_id, time)
            # print("URL LIST AT FIREBASE - ", file_url)
            for url in file_url:
                combined_text += process_file(url) + "\n\n***************************************************\n\n"       
                
            with open("output_final.txt", "w", encoding="utf-8") as final:
                final.write(combined_text)    
            
            print("Instructions - ", instruction)
            combined_summary = get_evaluation(combined_text, isInstruction=instruction)
            temp = markdown.markdown(combined_summary)
            session["result_generated"] = True 
            print("result marked true")
            return render_template("summary_out.html", output=temp)
        elif check_file:
            files = request.files.getlist('file')
            for f in files:
                if f and allowed_file(f.filename):
                    upload_files(f)
                else:
                    combined_text += f"Invalid or unsupported file: {f.filename}\n"
                    
            session_id = session.get('session_id')
            file_url = get_user_files(session_id, time)
            for url in file_url:
                combined_text += process_file(url) + "\n\n***************************************************\n\n"
                
            with open("output_final.txt", "w", encoding="utf-8") as final:
                final.write(combined_text)
            
            combined_summary = get_evaluation(combined_text)
            temp = markdown.markdown(combined_summary)
            session["result_generated"] = True 
            print("result marked true")
            return render_template("summary_out.html", output=temp)
        else:
            return render_template("summary_out.html", output="<p>Invalid input received.</p>")

# Route for deleting user files from firebase bucket
@app.route('/delete-user-files', methods=['GET'])
def accountcleanup():
    print("Delete Function is called.")
    expiration_time = datetime.now(timezone.utc) - timedelta(minutes=2)

    docs = db.collection("user_files").where(
        "uploaded_at", "<", expiration_time).stream()

    deleted_files = []

    for doc in docs:
        data = doc.to_dict()
        session_id = data.get("session_id")
        filename = data.get("filename")

        if session_id and filename:
            file_path = f"sessions/{session_id}/{filename}"
            blob = bucket.blob(file_path)

            try:
                blob.delete()
                db.collection("user_files").document(doc.id).delete()
                print(f"Deleted: {file_path}")
                deleted_files.append(filename)
            except Exception as e:
                print(f"Failed to delete {file_path}: {e}")

    if not deleted_files:
        return "No Files at Firebase\n", 200

    return jsonify({"deleted_files": deleted_files}), 200

# Route for evalation page
@app.route('/evaluate', methods=['GET', 'POST'])
def evaluate():
    try:
        if request.method == 'POST':
            if session.get("result_generated"):
                print("Rechecked result generation reload")
                return redirect(url_for("input"))
            if 'file' in request.files and request.files['file'].filename:
                file = request.files['file']
                if file.filename == '':
                    return "No file selected.", 400

                if file and allowed_file(file.filename):
                    time = datetime.now(timezone.utc)
                    upload_files(file)
                    session_id = session.get('session_id')
                    file_url = get_user_files(session_id, time)
                    file_content, name = read_file_from_gcs(file_url[0])
                    file_extension = name.split(".")[-1].lower()

                    if file_extension == 'txt':
                        response, score_response = process_txt_file(
                            file_content.decode('utf-8'))
                    elif file_extension == 'pdf':
                        response, score_response = process_pdf_file(
                            file_content)
                    elif file_extension in {'png', 'jpg', 'jpeg'}:
                        response, score_response = process_img_file(
                            file_content)
                    elif file_extension == 'docx':
                        response, score_response = process_docx_file(
                            file_content)
                    else:
                        return "File type not allowed", 400

                    score_list = score_response.text.split(":")
                    if len(score_list) > 1:
                        score = score_list[1].strip()
                        if not score.isdigit():
                            score = "-/"
                    else:
                        score = "-/"

                    evaluation_md = response.text
                    evaluation_html = markdown.markdown(evaluation_md)

                    session["result_generated"] = True 
                    print("result marked true")
                    return render_template('evaluate.html', evaluation=evaluation_html, score=score)

                else:
                    print("File not supported.")
                    session["result_generated"] = True 
                    print("result marked true")
                    return render_template('evaluate.html', evaluation="File format is not supported.", score="-/")
            elif 'fname' in request.form:
                text = request.form['fname']
                response, score_response = get_evaluate(text)
                score_list = score_response.text.split(":")
                if len(score_list) > 1:
                    score = score_list[1].strip()
                    if not score.isdigit():
                        score = "-/"
                else:
                    score = "-/"

                evaluation_md = response.text
                evaluation_html = markdown.markdown(evaluation_md)
                session["result_generated"] = True 
                print("result marked true")
                return render_template('evaluate.html', evaluation=evaluation_html, score=score)
            else:
                return "No input provided.", 400

    except Exception as e:
        return f"An error occurred: {str(e)}", 500

    return render_template("evaluate.html", prompt="Submit a file or text.", score="-/")


# Run the app in debug mode for development
if __name__ == '__main__':
    app.run(debug=True)

# End of code
