from flask import Flask, render_template, request, session
import os
import json
from werkzeug.utils import secure_filename
import google.generativeai as genai
import base64
import docx
from odfdo import Document
import markdown
import requests
import tempfile
from google.cloud import storage, vision
from datetime import datetime, timezone
from PIL import Image
from io import BytesIO
import uuid

with open('code/config.json', 'r') as c:
    params = json.load(c)["params"]

# Configure the Google Generative AI API
genai.configure(api_key=params['gen_api'])
model = genai.GenerativeModel("gemini-2.0-flash")

# Enhanced prompts for AI evaluation
NOTES_PROMPT = '''Provide detailed notes on the following topic:
- Include key concepts, principles, and subtopics .
- Provide hands-on exercises, projects, and challenges that reinforce the concepts learned.
- List the most effective books, courses, websites, and tools to master the topic. Each resource should be a clickable working link with a brief description (1-2 lines).
- Format the output in Markdown for easy readability. (Do not add any markdown word), max output is just 1000 so answer should be according to that limit.'''

SUMMARY_PROMPT = '''"Provide a concise summary of the given file, organized in bullet points and grouped by key topics. 
Focus solely on the summary, excluding any additional commentary or analysis. Use as many words as needed to accurately capture the content.
max token is 1000 so summarize the content accoridng to that ,(Do not add any markdown word)'''

# Initialize Flask app
app = Flask(__name__)
app.secret_key = os.urandom(24)

# Initialize GCS client
service_account_path = 'code/service-account.json'
storage_client = storage.Client.from_service_account_json(service_account_path)
bucket = storage_client.bucket(params['gcs_bucket_name'])

# Initialize Google Cloud Vision client
if 'GOOGLE_APPLICATION_CREDENTIALS' in os.environ:
    vision_client = vision.ImageAnnotatorClient()
else:
    vision_client = vision.ImageAnnotatorClient.from_service_account_json(service_account_path)

# Define allowed file extensions
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'txt', 'pdf', 'docx', 'odt'}

# Function to check allowed file extensions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Function to evaluate text content
def get_evaluation(text, is_file=False):
    if is_file:
        response = model.generate_content(
            [text, SUMMARY_PROMPT], generation_config=genai.GenerationConfig(
                max_output_tokens=1000,  # Set the token limit back to 1500
                temperature=0.5))
        content = response.candidates[0].content.parts[0].text if response.candidates else "No summary generated"
    else:
        response = model.generate_content(
            [text, NOTES_PROMPT], generation_config=genai.GenerationConfig(
                max_output_tokens=1000,  # Set the token limit back to 1500
                temperature=0.5))
        content = response.candidates[0].content.parts[0].text if response.candidates else "No response generated"
    
    return content

def process_txt_file(content):
    return get_evaluation(content, is_file=True)

# Function to process .pdf files
def process_pdf_file(content):
    response = model.generate_content(
        [{'mime_type': 'application/pdf', 'data': content}, SUMMARY_PROMPT], generation_config=genai.GenerationConfig(
            max_output_tokens=1500,
            temperature=0.5))
    return response.candidates[0].content.parts[0].text if response.candidates else "No summary generated"

# Function to process image files (png, jpg, jpeg)
def process_img_file(content):
    image = vision.Image(content=content)
    response = vision_client.text_detection(image=image)
    annotations = response.text_annotations
    if annotations:
        img_text = annotations[0].description
        return get_evaluation(img_text, is_file=True)
    else:
        raise ValueError("No text detected in the image")

# Function to process .docx files
def process_docx_file(content):
    doc = docx.Document(content)
    full_text = [paragraph.text for paragraph in doc.paragraphs]
    docx_text = '\n'.join(full_text)
    return get_evaluation(docx_text, is_file=True)

# Function to process .odt files
def process_odt_file(content):
    odt_file = Document(content)
    text_content = [para.text for para in odt_file.body.get_elements("//text:p")]
    odt_text = '\n'.join(text_content)
    return get_evaluation(odt_text, is_file=True)

# Function to upload file to Google Cloud Storage
def upload_files(file):
    session_id = session.get('session_id')
    unique_filename = f"{uuid.uuid4()}.{file.filename.split('.')[-1]}"
    folder_name = f"sessions/{session_id}/"
    blob = bucket.blob(f"{folder_name}{unique_filename}")
    blob.upload_from_file(file)

    print(session_id)
    print("file uploaded to GCS")

    return f"https://storage.googleapis.com/{bucket.name}/{folder_name}{unique_filename}"

# Function to get file content from Google Cloud Storage
def get_file_content_from_gcs(url):
    response = requests.get(url)
    response.raise_for_status()
    response.encoding = 'utf-8'
    return response.content

# Set session before each request
@app.before_request
def set_session():
    if 'session_id' not in session:
        session['session_id'] = str(uuid.uuid4())
        print("Session ID - ", session['session_id'])

# Route for home page to load summary.html
@app.route('/')
def home():
    return render_template("summary.html")

# Route for summary page
@app.route('/summary')
def summary():
    return render_template("summary.html")

@app.route('/summary_out', methods=['POST'])
def summary_out():
    if request.method == 'POST':
        check_file = 'file' in request.files and request.files['file'].filename
        check_fname = 'fname' in request.form and request.form['fname']
        if check_file:
            f = request.files['file']
            if f and allowed_file(f.filename):
                # Upload to GCS first
                public_url = upload_files(f)
                # Reset file pointer to beginning since upload_files consumed it
                f.seek(0)
                file_content = f.read()
                file_summary = process_file(file_content, f.filename)
                output = f"<h3>File Summary:</h3><div class='styled-content'>{markdown.markdown(file_summary)}</div>"
                return render_template("summary_out.html", output=output)
            else:
                return render_template("summary_out.html", output="<p>Invalid or unsupported file.</p>")
        elif check_fname:
            data = request.form['fname']
            text_notes = get_evaluation(data)
            output = f"<h3>Text Notes:</h3><div class='styled-content'>{markdown.markdown(text_notes)}</div>"
            return render_template("summary_out.html", output=output)
        else:
            return render_template("summary_out.html", output="<p>Invalid input received.</p>")

def process_file(content, filename):
    if filename.endswith('.txt'):
        return process_txt_file(content)
    elif filename.endswith('.pdf'):
        return process_pdf_file(content)
    elif filename.endswith('.png') or filename.endswith('.jpg') or filename.endswith('.jpeg'):
        return process_img_file(content)
    elif filename.endswith('.docx'):
        return process_docx_file(content)
    elif filename.endswith('.odt'):
        return process_odt_file(content)
    else:
        raise ValueError("Unsupported file type")

# Run the app in debug mode for development
if __name__ == '__main__':
    app.run(debug=True)
